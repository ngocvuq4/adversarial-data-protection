from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
ASSETS = REPORTS / "assets" / "benchmark_tables"
OUT_DOCX = REPORTS / "Unlearnable_Examples_Benchmark_Report.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: object, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(8)
    if color is not None:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], "24292F")
        set_cell_text(hdr[idx], header, bold=True, color=RGBColor(255, 255, 255))
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            if row_idx % 2 == 1:
                set_cell_shading(cells[col_idx], LIGHT_GRAY)
            set_cell_text(cells[col_idx], value)
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 99, 110)


def add_image(doc: Document, image_name: str, caption: str) -> None:
    image_path = ASSETS / image_name
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.5))
    add_caption(doc, caption)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    title = p.add_run("Báo Cáo Thực Nghiệm\nUnlearnable Examples")
    title.bold = True
    title.font.size = Pt(26)
    title.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = p.add_run("Adversarial Data Protection Framework")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = RGBColor(89, 99, 110)

    doc.add_paragraph()
    meta = [
        ("Project", "Adversarial Data Protection Framework"),
        ("Kỹ thuật trọng tâm", "Unlearnable Examples / Error-Minimizing Noise"),
        ("Dataset chính", "CIFAR-10"),
        ("Surrogate chính", "ResNet-18 đã điều chỉnh cho CIFAR-10"),
        ("Victim models", "ResNet-18, MobileNetV2, VGG-16"),
        ("Kết quả chính", "Protected clean test accuracy giảm về gần random guessing (~0.10)"),
    ]
    add_table(doc, ["Mục", "Thông tin"], meta)
    doc.add_page_break()


def build_doc() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title_page(doc)

    doc.add_heading("1. Tổng Quan Project", level=1)
    doc.add_paragraph(
        "Project này nghiên cứu nhóm kỹ thuật làm nhiễu dữ liệu ảnh nhằm ngăn chặn việc trích xuất đặc trưng "
        "phục vụ huấn luyện mô hình học sâu. Trọng tâm hiện tại là kỹ thuật Unlearnable Examples, trong đó "
        "nhiễu được tối ưu để mô hình học các shortcut giả thay vì học đặc trưng thật của ảnh."
    )
    doc.add_paragraph(
        "Mục tiêu thực nghiệm không phải là làm mô hình phân loại sai trên chính ảnh đã nhiễu, mà là kiểm tra "
        "xem mô hình sau khi được huấn luyện trên protected dataset có còn tổng quát hóa được lên clean test set hay không."
    )

    doc.add_heading("2. Cơ Chế Unlearnable Examples", level=1)
    doc.add_paragraph(
        "Unlearnable Examples trong project này được triển khai theo hướng error-minimizing noise. Với mỗi lớp ảnh, "
        "hệ thống học một nhiễu class-wise có chuẩn L-infinity bị giới hạn bởi epsilon. Sau đó nhiễu này được cộng "
        "vào ảnh clean để tạo protected dataset."
    )
    doc.add_paragraph("Quy trình sinh nhiễu có hai vòng lặp chính:")
    add_bullet(doc, "Vòng trong: huấn luyện surrogate model tạm thời trên dữ liệu đã được thêm nhiễu.")
    add_bullet(doc, "Vòng ngoài: cập nhật noise bằng PGD dựa trên gradient theo noise/input, không cập nhật trực tiếp weight của victim.")
    doc.add_paragraph(
        "Cách làm này khiến mô hình dễ học tín hiệu nhiễu hơn đặc trưng thật. Khi đánh giá trên clean test set, "
        "những shortcut đó không còn tồn tại, dẫn đến accuracy giảm mạnh."
    )

    doc.add_heading("3. Thuật Ngữ Trong Thực Nghiệm", level=1)
    add_table(
        doc,
        ["Thuật ngữ", "Ý nghĩa trong project"],
        [
            ["Clean dataset", "Dữ liệu gốc chưa thêm nhiễu."],
            ["Protected dataset", "Dữ liệu training sau khi áp dụng unlearnable noise."],
            ["Surrogate model", "Mô hình dùng để sinh noise. Trong thí nghiệm chính là ResNet-18."],
            ["Victim model", "Mô hình được train trên protected dataset để kiểm tra hiệu quả bảo vệ."],
            ["Clean test accuracy", "Accuracy của victim trên test set sạch, dùng để đo mô hình có học thật hay không."],
            ["ASR proxy", "Attack success proxy, tính gần đúng bằng 1 - protected clean test accuracy."],
        ],
    )

    doc.add_heading("4. Protocol Đánh Giá", level=1)
    doc.add_paragraph("Protocol đánh giá được chuẩn hóa như sau:")
    for step in [
        "Train baseline victim trên clean training set.",
        "Đánh giá baseline victim trên clean test set.",
        "Sinh class-wise unlearnable noise bằng surrogate model.",
        "Áp dụng noise lên clean training set để tạo protected dataset.",
        "Train victim model mới trên protected dataset.",
        "Đánh giá victim này trên clean test set.",
    ]:
        add_number(doc, step)
    doc.add_paragraph(
        "Điểm quan trọng nhất là victim sau khi train trên protected dataset phải được đánh giá trên clean test set. "
        "Nếu đánh giá trên ảnh đã nhiễu, kết quả sẽ không phản ánh việc mô hình có học được đặc trưng thật hay không."
    )

    doc.add_heading("5. Giai Đoạn 1: Thử Nghiệm Ban Đầu Trên Subset 5000", level=1)
    doc.add_paragraph(
        "Giai đoạn đầu dùng subset 5000 ảnh để kiểm tra pipeline sinh nhiễu, train baseline, train victim trên protected dataset "
        "và tính các metric cơ bản. Ở giai đoạn này chưa có validation protocol chặt chẽ, nên kết quả chỉ dùng để thăm dò."
    )
    add_table(
        doc,
        ["Epsilon", "Train size", "Victim", "Baseline acc", "Protected acc", "Drop", "ASR", "PSNR", "SSIM"],
        [
            [0.01, 5000, "ResNet-18", 0.3690, 0.1160, 0.2530, 0.8840, 40.6858, 0.9928],
            [0.03, 5000, "ResNet-18", 0.4900, 0.1230, 0.3670, 0.8770, 31.4192, 0.9413],
            [0.05, 5000, "ResNet-18", 0.3850, 0.1170, 0.2680, 0.8830, 27.1187, 0.8834],
        ],
    )
    add_image(doc, "benchmark_pre_validation_no_seed.png", "Bảng kết quả thử nghiệm ban đầu khi chưa cố định seed.")
    doc.add_paragraph(
        "Nhận xét: protected accuracy đã giảm về gần 0.10, nhưng baseline accuracy dao động giữa các lần chạy. "
        "Điều này cho thấy pipeline có hiệu ứng, nhưng chưa đủ chặt để dùng làm kết quả chính."
    )

    doc.add_heading("6. Giai Đoạn 2: Cố Định Seed", level=1)
    doc.add_paragraph(
        "Sau đó, seed được cố định bằng giá trị 42 để giảm sai khác giữa các lần chạy. Baseline accuracy trở nên thống nhất hơn, "
        "giúp so sánh epsilon công bằng hơn."
    )
    add_table(
        doc,
        ["Epsilon", "Seed", "Baseline acc", "Protected acc", "Drop", "ASR", "PSNR", "SSIM"],
        [
            [0.01, 42, 0.4180, 0.1000, 0.3180, 0.9000, 40.7485, 0.9918],
            [0.03, 42, 0.4180, 0.0850, 0.3330, 0.9150, 31.2963, 0.9424],
            [0.05, 42, 0.4180, 0.1050, 0.3130, 0.8950, 27.0640, 0.8818],
        ],
    )
    add_image(doc, "benchmark_pre_validation_seed42.png", "Bảng kết quả subset 5000 sau khi cố định seed.")
    doc.add_paragraph(
        "Nhận xét: epsilon 0.03 cho protected accuracy thấp nhất trong nhóm này. Tuy nhiên baseline chỉ đạt 0.4180, "
        "vẫn còn thấp nên chưa nên dùng làm benchmark cuối."
    )

    doc.add_heading("7. Giai Đoạn 3: Thêm Validation Để Chọn Epsilon", level=1)
    doc.add_paragraph(
        "Giai đoạn tiếp theo bổ sung validation split. Mục đích là không chọn epsilon trực tiếp theo test set. "
        "Cấu hình gồm 10000 ảnh train, 2000 ảnh validation và 10000 ảnh test."
    )
    add_table(
        doc,
        ["Eps", "Train", "Val", "Test", "Baseline val", "Protected val", "Val ASR", "PSNR", "SSIM", "Quality OK"],
        [
            [0.01, 10000, 2000, 10000, 0.6510, 0.1375, 0.8625, 40.5664, 0.9922, "True"],
            [0.03, 10000, 2000, 10000, 0.6510, 0.1055, 0.8945, 31.1111, 0.9348, "True"],
            [0.05, 10000, 2000, 10000, 0.6510, 0.1030, 0.8970, 26.6897, 0.8619, "False"],
        ],
    )
    add_image(doc, "benchmark_validation_epsilon_ablation.png", "Validation ablation để chọn epsilon.")
    add_image(doc, "unlearnable_validation_epsilon_line_chart.png", "Biểu đồ trade-off giữa attack strength và image quality.")
    doc.add_paragraph(
        "Epsilon 0.05 làm protected validation accuracy thấp, nhưng SSIM giảm xuống 0.8619, nghĩa là chất lượng ảnh suy giảm rõ hơn. "
        "Epsilon 0.01 giữ ảnh rất tốt nhưng attack yếu hơn. Epsilon 0.03 là điểm cân bằng tốt giữa hiệu quả bảo vệ và chất lượng ảnh."
    )
    add_table(
        doc,
        ["Selected epsilon", "Baseline test acc", "Protected test acc", "Final drop", "Final ASR", "PSNR", "SSIM"],
        [[0.03, 0.6457, 0.1012, 0.5445, 0.8988, 31.1111, 0.9348]],
    )
    add_image(doc, "benchmark_validation_final_test.png", "Kết quả final test sau khi chọn epsilon bằng validation.")

    doc.add_heading("8. Giai Đoạn 4: Full Final Benchmark", level=1)
    doc.add_paragraph(
        "Sau khi chọn được epsilon 0.03, thí nghiệm final được chạy trên split lớn hơn: 45000 ảnh train, "
        "5000 ảnh validation và 10000 ảnh test. Đây là kết quả chính nên dùng trong báo cáo."
    )
    add_table(
        doc,
        ["Eps", "Train", "Val", "Test", "Victim", "Baseline test", "Protected val", "Val ASR", "PSNR", "SSIM"],
        [[0.03, 45000, 5000, 10000, "ResNet-18", 0.8528, 0.0966, 0.9034, 31.5549, 0.9419]],
    )
    add_image(doc, "benchmark_full_final.png", "Full-final benchmark trên split 45000/5000/10000.")
    add_table(
        doc,
        ["Selected epsilon", "Baseline test acc", "Protected test acc", "Test drop", "Test ASR", "Runtime"],
        [[0.03, 0.8528, 0.1037, 0.7491, 0.8963, "3121.89s"]],
    )
    add_image(doc, "benchmark_full_final_test.png", "Kết quả final test của ResNet-18 victim.")
    doc.add_paragraph(
        "Baseline ResNet-18 đạt 0.8528 clean test accuracy khi train trên clean data. Khi train trên protected dataset, "
        "clean test accuracy chỉ còn 0.1037, gần mức đoán ngẫu nhiên của CIFAR-10. Đây là bằng chứng mạnh cho thấy "
        "protected dataset làm mô hình không học được đặc trưng tổng quát."
    )

    doc.add_heading("9. Giai Đoạn 5: Transfer Victims", level=1)
    doc.add_paragraph(
        "Ở giai đoạn transfer, protected dataset đã sinh bằng ResNet-18 surrogate được dùng lại để train các victim model khác. "
        "Mục tiêu là kiểm tra nhiễu có chỉ hoạt động với ResNet-18 hay có thể chuyển sang kiến trúc khác."
    )
    add_table(
        doc,
        ["Surrogate", "Victim", "Eps", "Train", "Test", "Baseline acc", "Protected acc", "Drop", "ASR"],
        [
            ["ResNet-18", "MobileNetV2", 0.03, 45000, 10000, 0.6570, 0.1008, 0.5562, 0.8992],
            ["ResNet-18", "VGG-16", 0.03, 45000, 10000, 0.8291, 0.1000, 0.7291, 0.9000],
        ],
    )
    add_image(doc, "benchmark_transfer_victims.png", "Benchmark transferability trên MobileNetV2 và VGG-16.")
    doc.add_paragraph(
        "Kết quả cho thấy cả MobileNetV2 và VGG-16 đều rơi về gần 0.10 clean test accuracy sau khi train trên protected dataset. "
        "Điều này chứng minh noise có khả năng transfer sang kiến trúc khác, không chỉ giới hạn ở model surrogate."
    )

    doc.add_heading("10. Ý Nghĩa Các Chỉ Số", level=1)
    add_table(
        doc,
        ["Chỉ số", "Ý nghĩa"],
        [
            ["Baseline clean test acc", "Accuracy của model train trên clean data và test trên clean test set."],
            ["Protected clean test acc", "Accuracy của model train trên protected data nhưng test trên clean test set."],
            ["Drop", "Mức giảm accuracy giữa baseline và protected victim."],
            ["ASR proxy", "Xấp xỉ mức thành công của attack, thường tính gần bằng 1 - protected clean test accuracy."],
            ["PSNR", "Đo mức sai khác ảnh theo tín hiệu/nhiễu; càng cao ảnh càng giống gốc."],
            ["SSIM", "Đo độ tương đồng cấu trúc ảnh; càng gần 1 ảnh càng giống gốc."],
            ["L-infinity", "Giới hạn biên độ nhiễu tối đa trên mỗi pixel."],
        ],
    )

    doc.add_heading("11. Kết Luận", level=1)
    doc.add_paragraph(
        "Qua các giai đoạn thử nghiệm, project đã xây dựng được một pipeline Unlearnable Examples hoàn chỉnh: "
        "sinh nhiễu bằng PGD, tạo protected dataset, train victim trên protected dataset, đánh giá trên clean test set, "
        "chọn epsilon bằng validation và kiểm tra transferability trên nhiều kiến trúc."
    )
    doc.add_paragraph(
        "Kết quả chính cho thấy ResNet-18 victim giảm từ 0.8528 xuống 0.1037 clean test accuracy. "
        "MobileNetV2 và VGG-16 cũng giảm về khoảng 0.10 khi train trên cùng protected dataset. "
        "Điều này phù hợp với mục tiêu đề tài: làm nhiễu dữ liệu nhằm ngăn chặn việc trích xuất đặc trưng phục vụ huấn luyện mô hình học sâu."
    )

    doc.add_heading("12. Giới Hạn Và Hướng Tiếp Theo", level=1)
    add_bullet(doc, "Kết quả hiện tại tập trung vào CIFAR-10, ảnh 32x32 và bài toán image classification.")
    add_bullet(doc, "Unlearnable hiện dùng class-wise noise; sample-wise noise có thể là hướng mở rộng.")
    add_bullet(doc, "Cloaking và CLIP-space Concept Poisoning cần benchmark riêng với metric khớp mục tiêu từng kỹ thuật.")
    add_bullet(doc, "Với ảnh độ phân giải cao, cần đánh giá thêm pipeline patch-based hoặc resize-based trong demo thực tế.")

    doc.add_heading("Phụ Lục: Notebook Và Kết Quả", level=1)
    add_table(
        doc,
        ["Notebook", "Vai trò"],
        [
            ["colab_unlearnable_experiment.ipynb", "Thử nghiệm ban đầu trên subset 5000."],
            ["colab_unlearnable_validation_experiment.ipynb", "Validation ablation để chọn epsilon."],
            ["colab_unlearnable_full_final_experiment.ipynb", "Chạy benchmark chính trên split 45000/5000/10000."],
            ["colab_unlearnable_transfer_victims.ipynb", "Kiểm tra transferability trên MobileNetV2 và VGG-16."],
        ],
    )

    doc.save(OUT_DOCX)
    print(f"Saved {OUT_DOCX}")


if __name__ == "__main__":
    build_doc()
